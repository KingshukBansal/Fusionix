from flask import Flask, request, jsonify
from docx import Document
from pymongo import MongoClient
import gridfs
from bson import ObjectId
from bson.errors import InvalidId

app = Flask(__name__)

# Connect to MongoDB using pymongo (you need to have pymongo installed)
client = MongoClient("mongodb://host.minikube.internal:27017/docx_files")  # Connect to the MongoDB instance
db = client['docx_files']  # Use the database where the DOCX files are stored
fs = gridfs.GridFS(db)  # Initialize GridFS to manage large file storage

# Route to fetch metadata of a DOCX file based on fid
@app.route('/metadata', methods=['GET'])
def get_metadata():
    fid = request.args.get("fid")  # Get the fid from the query string

    if not fid:
        return jsonify({"error": "fid is required"}), 400

    try:
        # Validate ObjectId
        try:
            file_id = ObjectId(fid)
        except InvalidId:
            return jsonify({"error": "Invalid fid format"}), 400
        
        # Retrieve the DOCX file from MongoDB GridFS using the fid
        docx_file = fs.get(file_id)

        # Load the DOCX file using python-docx
        document = Document(docx_file)

        # Extract metadata
        metadata = {
            'title': document.core_properties.title if document.core_properties.title else None,
            'author': document.core_properties.author if document.core_properties.author else None,
            'subject': document.core_properties.subject if document.core_properties.subject else None,
            'keywords': document.core_properties.keywords if document.core_properties.keywords else None,
            'created': document.core_properties.created if document.core_properties.created else None,
            'modified': document.core_properties.modified if document.core_properties.modified else None,
            'word_count': sum(len(paragraph.text.split()) for paragraph in document.paragraphs),
            'paragraph_count': len(document.paragraphs)
        }

        # Return metadata as JSON response
        return jsonify(metadata)

    except InvalidId:
        return jsonify({"error": "Invalid fid format"}), 400
    except gridfs.errors.NoFile:
        return jsonify({"error": "File not found"}), 404
    except Exception as e:
        return jsonify({"error": f"Error fetching file or extracting metadata: {str(e)}"}), 500


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080)
